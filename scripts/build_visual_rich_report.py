from __future__ import annotations

import hashlib
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submission" / "report" / "2026 오픈소스 개발자대회 결과보고서_1123(AI쀼).docx"
OUTPUT = ROOT / "submission" / "report" / "2026 오픈소스 개발자대회 결과보고서_1123(AI쀼)_시각자료보강본.docx"
VISUALS = ROOT / "submission" / "report" / "visuals"
EXPECTED_SOURCE_SHA256 = "9454a546bbca7b50028a2d89026a8e491af83d1b103a8e3cc9b4c720b605988e"


FIGURES = [
    (
        VISUALS / "01-system-architecture.png",
        "그림 1. MCP 클라이언트 요청부터 네 영역 검사, 근거 생성, SHIP/HOLD 판정까지의 핵심 흐름",
        "K-Guard MCP 시스템 구성도",
    ),
    (
        VISUALS / "02-four-domain-feature-map.png",
        "그림 2. 사이트·API·데이터·운영 영역을 하나의 Guardian 판정으로 결합한 검수 범위",
        "K-Guard MCP 네 영역 기능 지도",
    ),
    (
        VISUALS / "03-review-lifecycle-and-runtime-block.png",
        "그림 3. 코드 변경 시 이전 판정을 무효화하고 재검수하며, 런타임 공격은 403과 감사 영수증으로 연결",
        "K-Guard MCP 검수 생명주기와 런타임 차단",
    ),
    (
        VISUALS / "04-ai-development-roles.png",
        "그림 4. 구현·읽기 전용 감사·통합 검증을 분리하고 사람이 최종 책임을 지는 AI 협업 구조",
        "K-Guard MCP 개발 과정의 AI 역할 분리",
    ),
    (
        VISUALS / "05-expected-impact-and-use-cases.png",
        "그림 5. 전문 보안인력이 없는 개발 현장과 개인정보 서비스에 적용할 수 있는 기대효과와 활용 분야",
        "K-Guard MCP 기대효과와 활용 분야",
    ),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def iter_paragraphs(parent):
    for paragraph in getattr(parent, "paragraphs", []):
        yield paragraph
    for table in getattr(parent, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def set_run_font(run, size: float, bold: bool = False, color: str = "202827") -> None:
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Malgun Gothic")
    r_fonts.set(qn("w:hAnsi"), "Malgun Gothic")
    r_fonts.set(qn("w:eastAsia"), "맑은 고딕")


def move_before(paragraph, boundary) -> None:
    paragraph._p.getparent().remove(paragraph._p)
    boundary.addprevious(paragraph._p)


def new_paragraph_before(document: Document, boundary):
    paragraph = document.add_paragraph()
    move_before(paragraph, boundary)
    return paragraph


def add_page_break(document: Document, boundary) -> None:
    paragraph = new_paragraph_before(document, boundary)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_title(document: Document, boundary, title: str, subtitle: str) -> None:
    paragraph = new_paragraph_before(document, boundary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(title), 14.5, bold=True, color="173C2F")

    paragraph = new_paragraph_before(document, boundary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(subtitle), 8.5, color="4C5C57")


def add_figure(document: Document, boundary, image_path: Path, caption: str, alt_text: str) -> None:
    paragraph = new_paragraph_before(document, boundary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(5.82))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", alt_text)
    doc_pr.set("descr", alt_text)

    paragraph = new_paragraph_before(document, boundary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = False
    set_run_font(paragraph.add_run(caption), 7.6, color="56615D")


def add_footer_note(document: Document, boundary, text: str) -> None:
    paragraph = new_paragraph_before(document, boundary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(text), 7.5, bold=True, color="2E6B55")


def build() -> dict:
    source_sha = sha256(SOURCE)
    if source_sha != EXPECTED_SOURCE_SHA256:
        raise RuntimeError(f"Source report changed: {source_sha}")
    missing = [str(path) for path, _, _ in FIGURES if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing figures: " + ", ".join(missing))

    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    # The narration/video was finalized at exactly three minutes after this
    # report draft was written. Keep the public-facing report synchronized.
    replacements = 0
    for paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            if "168초" in run.text:
                run.text = run.text.replace("168초", "180초")
                replacements += 1

    body = document._element.body
    section_break_paragraph = next(
        child for child in body if child.tag == qn("w:p") and child.xpath(".//w:sectPr")
    )

    add_figure(document, section_break_paragraph, *FIGURES[4])
    add_page_break(document, section_break_paragraph)
    add_title(
        document,
        section_break_paragraph,
        "시각자료 1 · 시스템 구성과 검수 범위",
        "AI 코딩 흐름 안에 검수와 재검수를 붙이고, 결과를 출하 판단까지 연결합니다.",
    )
    add_figure(document, section_break_paragraph, *FIGURES[0])
    add_figure(document, section_break_paragraph, *FIGURES[1])

    add_page_break(document, section_break_paragraph)
    add_title(
        document,
        section_break_paragraph,
        "시각자료 2 · 실제 검수 순환과 AI 협업",
        "한 번의 탐지로 끝내지 않고 수정·재검수·런타임 차단·감사 근거까지 이어갑니다.",
    )
    add_figure(document, section_break_paragraph, *FIGURES[2])
    add_figure(document, section_break_paragraph, *FIGURES[3])
    add_footer_note(document, section_break_paragraph, "운영 원칙 · AI는 개발과 검증을 보조하며 최종 출하 책임은 사람에게 있습니다.")

    document.save(OUTPUT)
    reopened = Document(OUTPUT)
    result = {
        "source": str(SOURCE),
        "source_sha256": source_sha,
        "output": str(OUTPUT),
        "output_sha256": sha256(OUTPUT),
        "embedded_figures": len(reopened.inline_shapes),
        "sections": len(reopened.sections),
        "duration_text_replacements": replacements,
        "standalone_visual_assets": 6,
    }
    if result["embedded_figures"] != 5 or result["sections"] != 2:
        raise RuntimeError(f"Unexpected output structure: {result}")
    return result


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
